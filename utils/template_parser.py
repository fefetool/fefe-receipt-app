from docx import Document
from io import BytesIO
import re
from difflib import get_close_matches

def analyze_template(template_file):
    doc = Document(BytesIO(template_file))
    result = {
        'tables': [],
        'paragraphs': [],
        'available_fields': [],
        'field_suggestions': {}
    }

    for i, para in enumerate(doc.paragraphs):
        para_info = {
            'index': i,
            'text': para.text.strip(),
            'style': para.style.name,
            'is_date_field': bool(re.search(r'年\s*月\s*日', para.text)),
            'is_title': any(x in para.text for x in ['收入憑證', '支出憑證'])
        }
        result['paragraphs'].append(para_info)

    for table_idx, table in enumerate(doc.tables):
        if table.rows:
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            result['available_fields'].extend(header_cells)

    return result

def guess_excel_field(template_field, excel_columns):
    match = get_close_matches(template_field, excel_columns, n=1, cutoff=0.6)
    return match[0] if match else None


# -*- coding: utf-8 -*-
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime
import os
import zipfile
from tempfile import NamedTemporaryFile
from docx.shared import Inches
from PIL import Image

def init_session_state():
    if 'field_mappings' not in st.session_state:
        st.session_state.field_mappings = []
    if 'logo_position' not in st.session_state:
        st.session_state.logo_position = "top-right"
    if 'image_position' not in st.session_state:
        st.session_state.image_position = "below-center"

def add_image_to_doc(doc, image_file, position, is_logo=False):
    img_path = NamedTemporaryFile(delete=False, suffix=".png").name
    with open(img_path, 'wb') as f:
        f.write(image_file.getvalue())

    try:
        img = Image.open(img_path)
        width, height = img.size
        max_width = Inches(6.5) if not is_logo else Inches(1.5)
        ratio = min(max_width / width, 1.0)
        new_width = int(width * ratio)
        new_height = int(height * ratio)

        if is_logo:
            section = doc.sections[0]
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

            if "左" in position:
                paragraph.alignment = 0
            elif "右" in position:
                paragraph.alignment = 2

            run = paragraph.add_run()
            run.add_picture(img_path, width=Inches(1.5))
        else:
            paragraph = doc.add_paragraph()
            if "置中" in position:
                paragraph.alignment = 1
            elif "右" in position:
                paragraph.alignment = 2
            else:
                paragraph.alignment = 0

            run = paragraph.add_run()
            run.add_picture(img_path, width=Inches(4.0))
    finally:
        os.unlink(img_path)

def main():
    st.set_page_config(page_title="智慧憑證生成系統", layout="wide")
    st.title("📄 智慧憑證生成系統")
    init_session_state()

    st.header("1. 文件上传")
    col1, col2 = st.columns(2)
    with col1:
        excel_file = st.file_uploader("📊 上传Excel收支明细", type=["xlsx"], key="excel_uploader")
    with col2:
        word_template = st.file_uploader("📄 上传Word凭证模板", type=["docx"], key="word_uploader")

    st.header("2. 日期格式设定")
    date_format = st.selectbox(
        "选择日期显示格式：",
        ["西元年 (2025年4月1日)", "民国年 (民国114年4月1日)"],
        key="date_format"
    )
# ...略，其餘保留在 canvas 中的程式碼...
